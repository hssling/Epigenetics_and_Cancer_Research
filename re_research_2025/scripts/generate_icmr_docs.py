from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import datetime

# Output Directory
OUTPUT_DIR = r"d:\research-automation\Epigenetics research\re_research_2025\submission_files"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Manuscript Details
TITLE = "Factors Influencing Epigenetics in Cancer Prevention: A Systematic Analysis of Recent Evidence (2024-2025)"
RUNNING_TITLE = "Epigenetics in Cancer Prevention 2024-2025"
AUTHORS = ["Research Team"] # Placeholder
WORD_COUNT = "Approx. 2000"
TABLES_COUNT = "1"
FIGURES_COUNT = "2"
DATE = datetime.datetime.now().strftime("%B %d, %Y")

def setup_document():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    return doc

def save_document(doc, filename):
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    print(f"Created: {path}")

# 1. First Page (Title Page)
def create_first_page():
    doc = setup_document()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph() # Spacer
    
    # Type of Article
    p = doc.add_paragraph("Type of Article: Systematic Review / Meta-Analysis")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Running Title
    doc.add_paragraph(f"Running Title: {RUNNING_TITLE}")
    
    doc.add_paragraph()
    
    # Authors
    doc.add_paragraph("Authors:")
    for auth in AUTHORS:
        doc.add_paragraph(auth)
        
    doc.add_paragraph()
    
    # Word Counts etc
    doc.add_paragraph(f"Word Count: {WORD_COUNT}")
    doc.add_paragraph(f"Number of Tables: {TABLES_COUNT}")
    doc.add_paragraph(f"Number of Figures: {FIGURES_COUNT}")
    doc.add_paragraph(f"Number of References: 8")
    
    doc.add_paragraph()
    doc.add_paragraph("Conflict of Interest: None declared.")
    doc.add_paragraph("Source of Support: None.")
    
    save_document(doc, "ijmr_first_page.docx")

# 2. Undertaking by Authors
def create_undertaking():
    doc = setup_document()
    doc.add_heading('UNDERTAKING BY AUTHORS', 0)
    
    text = (
        f"We, the undersigned, give an undertaking that the manuscript entitled \"{TITLE}\" submitted to the Indian Journal of Medical Research is original, has not been published, and is not currently under consideration for publication elsewhere.\n\n"
        "We agree to transfer all copyright ownership, including any and all rights incidental thereto, exclusively to the Indian Journal of Medical Research, in the event that such work is published by the Indian Journal of Medical Research.\n\n"
        "We state that the manuscript contains no violation of any existing copyright or other third party right or any material of an obscene, indecent, libellous or otherwise unlawful nature and that to the best of our knowledge and belief the manuscript does not infringe the rights of others.\n\n"
        f"Date: {DATE}\n\n"
        "Signatures:\n\n"
        "__________________________\n"
        "(Corresponding Author on behalf of all authors)"
    )
    doc.add_paragraph(text)
    save_document(doc, "Undertaking by Authors.docx")

# 3. Conflict of Interest
def create_coi():
    doc = setup_document()
    doc.add_heading('CONFLICT OF INTEREST STATEMENT', 0)
    
    text = (
        f"Manuscript Title: {TITLE}\n\n"
        "The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.\n\n"
        "There are no financial conflicts of interest to disclose.\n\n"
        "Sincerely,\n\n"
        "The Authors"
    )
    doc.add_paragraph(text)
    save_document(doc, "ijmr_conflict_of_interest.docx")

# 4. Copyright Transfer
def create_copyright():
    doc = setup_document()
    doc.add_heading('COPYRIGHT TRANSFER AGREEMENT', 0)
    
    text = (
        "To: The Editor-in-Chief,\n"
        "Indian Journal of Medical Research\n\n"
        f"Title of Article: {TITLE}\n\n"
        "I/We hereby assign and transfer to the Indian Council of Medical Research (ICMR) all rights of copyright and ownership of the above titled manuscript.\n\n"
        "I/We warrant that the article is original, does not infringe upon any copyright or other proprietary right of any third party, is not under consideration by another journal, and has not been previously published.\n\n"
        f"Date: {DATE}\n\n"
        "Signature(s) of Author(s):\n\n"
        "__________________________"
    )
    doc.add_paragraph(text)
    save_document(doc, "ijmr_copyright_transfer.docx")

# 5. Ethics Statement
def create_ethics():
    doc = setup_document()
    doc.add_heading('ETHICAL STATEMENT', 0)
    
    text = (
        f"Manuscript Title: {TITLE}\n\n"
        "This study is a systematic analysis of existing published literature openly available in the PubMed database. It involves the synthesis of secondary data and does not involve any direct interaction with human participants or animals. "
        "Therefore, formal ethical approval from an Institutional Review Board (IRB) or Ethics Committee was not required for this research.\n\n"
        "All data sources have been properly cited in accordance with academic standards."
    )
    doc.add_paragraph(text)
    save_document(doc, "ijmr_ethics_statement.docx")

if __name__ == "__main__":
    create_first_page()
    create_undertaking()
    create_coi()
    create_copyright()
    create_ethics()
