import csv
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# File Paths
input_md = r"d:\research-automation\Epigenetics research\re_research_2025\manuscript_v2.md"
input_csv = r"d:\research-automation\Epigenetics research\re_research_2025\data\extracted_data.csv"
output_docx = r"d:\research-automation\Epigenetics research\re_research_2025\manuscript_submission_packaged.docx"
assets_dir = r"d:\research-automation\Epigenetics research\re_research_2025\assets"

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def clean_markdown_for_docx(text):
    # Remove image links
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    # Remove HTML comments or tags if any
    text = re.sub(r'<.*?>', '', text)
    return text.strip()

def process_text_with_superscripts(paragraph, text):
    # Regex to find citations like [1] or [1, 2]
    patterns = r'(\[\d+(?:,\s*\d+)*\])'
    parts = re.split(patterns, text)
    
    for part in parts:
        if re.match(patterns, part):
            cit_nums = part.strip('[]')
            run = paragraph.add_run(cit_nums)
            run.font.superscript = True
        else:
            bold_parts = re.split(r'(\*\*.+?\*\*)', part)
            for subpart in bold_parts:
                if subpart.startswith('**') and subpart.endswith('**'):
                    run = paragraph.add_run(subpart.strip('*'))
                    run.bold = True
                else:
                    italic_parts = re.split(r'(\*.+?\*)', subpart)
                    for subsubpart in italic_parts:
                        if subsubpart.startswith('*') and subsubpart.endswith('*'):
                            run = paragraph.add_run(subsubpart.strip('*'))
                            run.italic = True
                        else:
                            paragraph.add_run(subsubpart)

def main():
    if os.path.exists(output_docx):
        try:
            os.remove(output_docx)
        except PermissionError:
            print(f"ERROR: Cannot delete existing file {output_docx}. It is open in another program.")
            return

    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Read CSV Data for Table
    studies = []
    with open(input_csv, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            studies.append(row)
            
    # Read Markdown
    with open(input_md, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Parse Title
    title_match = re.search(r'^#\s+(.+)$', md_content, re.MULTILINE)
    title_text = title_match.group(1) if title_match else "Manuscript"
    
    # Add Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    # Add Abstract
    # Extract Abstract Content
    abstract_start = md_content.find('**Abstract**')
    if abstract_start == -1:
         abstract_start = md_content.find('## Abstract')
         
    intro_start = md_content.find('## 1. Introduction')
    
    if abstract_start != -1 and intro_start != -1:
        abstract_text = md_content[abstract_start:intro_start].strip()
        abstract_text = re.sub(r'^(\*\*Abstract\*\*|## Abstract)', '', abstract_text).strip()
        
        doc.add_heading('Abstract', level=1)
        lines = abstract_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            p = doc.add_paragraph()
            process_text_with_superscripts(p, line)
            
    # Process Main Sections
    sections_pattern = r'##\s+(\d+\.\s+.*)'
    headers = list(re.finditer(sections_pattern, md_content))
    
    ref_start_index = -1
    
    for i, match in enumerate(headers):
        header_title = match.group(1)
        start_pos = match.end()
        end_pos = headers[i+1].start() if i+1 < len(headers) else len(md_content)
        
        if "References" in header_title:
            ref_start_index = match.start()
            start_pos = match.end()
            ref_content = md_content[start_pos:].strip()
            continue 

        section_body = md_content[start_pos:end_pos].strip()
        doc.add_heading(header_title, level=1)
        
        paragraphs = section_body.split('\n')
        for line in paragraphs:
            line = line.strip()
            if not line: continue
            
            if line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=2)
                continue
            
            if line.startswith('![') or 'Figures available in' in line:
                continue
                
            if line.startswith('```') or line.startswith('graph TD') or 'style ' in line:
                continue
            
            if "A total of 29 studies met the inclusion criteria" in line:
                line += " (see Table 1)."
            
            p = doc.add_paragraph()
            process_text_with_superscripts(p, line)

    # Add References
    if ref_start_index != -1:
        doc.add_page_break()
        doc.add_heading('6. References', level=1)
        
        ref_lines = ref_content.split('\n')
        for line in ref_lines:
            line = line.strip()
            if not line: continue
            if line.startswith('---'): continue
            if line.startswith('*Appendices'): continue
            
            if re.match(r'^\d+\.', line):
                p = doc.add_paragraph()
                p.add_run(line)
    
    # Add Tables
    doc.add_page_break()
    doc.add_heading('Tables', level=1)
    
    table_title = doc.add_paragraph("Table 1: Characteristics of Selected High-Relevance Studies")
    table_title.style = 'Caption'
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'PMID'
    hdr_cells[1].text = 'Year'
    hdr_cells[2].text = 'Intervention'
    hdr_cells[3].text = 'Cancer'
    hdr_cells[4].text = 'Marker'
    
    for study in studies[:20]:
        row_cells = table.add_row().cells
        row_cells[0].text = study.get('PMID', '')
        row_cells[1].text = study.get('Year', '')
        row_cells[2].text = study.get('Intervention', '')
        row_cells[3].text = study.get('Cancer Type', '')
        row_cells[4].text = study.get('Epigenetic Marker', '')

    # Add Figures
    doc.add_page_break()
    doc.add_heading('Figures', level=1)
    
    doc.add_paragraph("Figure 1: Distribution of Intervention Types in Epigenetics Research (2024-2025)").style = 'Caption'
    try:
        doc.add_picture(f"{assets_dir}\\Figure_1_Interventions.png", width=Inches(6))
    except Exception as e:
        doc.add_paragraph(f"[Error loading Figure 1: {e}]")

    doc.add_paragraph("\n")

    doc.add_paragraph("Figure 2: Frequency of Cancer Types Investigated").style = 'Caption'
    try:
        doc.add_picture(f"{assets_dir}\\Figure_2_Cancer_Types.png", width=Inches(6))
    except Exception as e:
        doc.add_paragraph(f"[Error loading Figure 2: {e}]")
        
    doc.add_paragraph("\n")
    
    # Figure 3
    doc.add_paragraph("Figure 3: Conceptual DAG of Epigenetic Modulation in Cancer Prevention").style = 'Caption'
    try:
        doc.add_picture(f"{assets_dir}\\Figure_3_DAG.png", width=Inches(6))
    except Exception as e:
        doc.add_paragraph(f"[Error loading Figure 3: {e}]")
    
    try:
        doc.save(output_docx)
        print(f"Document saved to {output_docx}")
    except PermissionError:
        print(f"Error: Could not save document. Please close {output_docx} if it is open.")
    except Exception as e:
        print(f"Unexpected error saving document: {e}")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        import traceback
        traceback.print_exc()
